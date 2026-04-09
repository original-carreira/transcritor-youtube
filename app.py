# ==============================
# IMPORTS
# ==============================
from flask import Flask, render_template, request, send_file, redirect, url_for
from app.services.transcription_service import TranscriptionService # Import da classe

from docx import Document
from docx.shared import Pt
import io
import webbrowser
import threading
import time
import socket

# ==============================
# CONFIGURAÇÃO DA APLICAÇÃO
# ==============================

app = Flask(__name__)
service = TranscriptionService() # Instância única para o app

# ==============================
# ROTA PRINCIPAL
# ==============================

@app.route('/', methods=['GET', 'POST'])
def index():
    url = None
    titulo = None
    transcricao = None
    thumbnail = None

    if request.method == 'POST':
        url = request.form.get('url', '').strip()

        if not url:
            transcricao = "Por favor, insira uma URL do YouTube."
        else:
            # CENTRALIZADO: O service.process agora resolve tudo (cache ou nova extração)
            resultado = service.process(url)

            # Se o retorno for o dicionário do cache ou objeto completo
            if isinstance(resultado, dict):
                transcricao = resultado.get("transcricao")
                titulo = resultado.get("titulo")
                thumbnail = resultado.get("thumbnail")
            else:
                # Fallback caso o service retorne apenas a string da transcrição
                transcricao = resultado
                titulo = service.obter_titulo_video(url)
                video_id = service.extrair_id(url)
                thumbnail = service.obter_thumbnail(video_id)

    # CENTRALIZADO: Histórico via service
    historico = service.listar_historico()
    
    return render_template(
        'index.html',
        transcricao=transcricao,
        url=url,
        titulo=titulo,
        thumbnail=thumbnail,
        historico=historico
    )

# ==============================
# DOWNLOAD TXT
# ==============================

@app.route('/download_txt', methods=['POST'])
def download_txt():
    texto = request.form.get('texto', '').strip()
    titulo = request.form.get('titulo', '').strip()

    if not texto:
        return "Nenhum conteúdo para download.", 400

    # CENTRALIZADO: Limpeza de nome via service
    nome_arquivo = service.limpar_nome_arquivo(titulo or "transcricao")

    buffer = io.BytesIO()
    buffer.write(texto.encode('utf-8'))
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{nome_arquivo}.txt",
        mimetype="text/plain; charset=utf-8"
    )

# ==============================
# DOWNLOAD DOCX
# ==============================

@app.route('/download_docx', methods=['POST'])
def download_docx():
    texto = request.form.get('texto', '').strip()
    titulo = request.form.get('titulo', '').strip()

    if not texto:
        return "Nenhum conteúdo para download.", 400

    # CENTRALIZADO: Limpeza de nome via service
    nome_arquivo = service.limpar_nome_arquivo(titulo or "transcricao")

    document = Document()
    document.add_heading(titulo or "Transcrição", 0)

    for paragrafo in texto.split("\n\n"):
        if paragrafo.strip():
            p = document.add_paragraph(paragrafo.strip())
            p.paragraph_format.space_after = Pt(12)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{nome_arquivo}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ==============================
# LIMPAR HISTÓRICO
# ==============================

@app.route('/limpar_historico', methods=['POST'])
def limpar_historico():
    # CENTRALIZADO: Limpeza via service
    service.limpar_cache()
    return redirect(url_for('index'))

# ==============================
# EXECUÇÃO LOCAL (Mantido original)
# ==============================

def servidor_esta_pronto(host="127.0.0.1", port=5000):
    while True:
        try:
            with socket.create_connection((host, port), timeout=1):
                return True
        except OSError:
            time.sleep(0.3)

def abrir_navegador():
    servidor_esta_pronto()
    webbrowser.open("http://127.0.0.1:5000")

if __name__ == '__main__':
    threading.Thread(target=abrir_navegador).start()
    app.run(debug=False)
