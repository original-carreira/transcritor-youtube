import io
import re
from docx import Document
from docx.shared import Pt


class DocxExporter:

    def export(self, data: str, titulo: str, nome_arquivo: str):
        document = Document()
        document.add_heading(titulo or "Transcrição", 0)

        # 🔥 normalização segura
        data = re.sub(r'\n{3,}', '\n\n', data.strip())

        for paragrafo in data.split("\n\n"):
            paragrafo = paragrafo.strip()

            if not paragrafo:
                continue

            p = document.add_paragraph(paragrafo)
            p.paragraph_format.space_after = Pt(12)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return {
            "buffer": buffer,
            "filename": f"{nome_arquivo}.docx",
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }